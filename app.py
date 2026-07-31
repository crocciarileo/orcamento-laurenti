import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime, timedelta
import io
import os
import uuid
import hashlib
from PIL import Image

# Exportações
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ReportLab para PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

# Configuração da página Streamlit
st.set_page_config(page_title="Orçamentos - Laurenti Móveis", page_icon="📝", layout="wide")

# -----------------------------------------------------------------------------
# 1. BANCO DE DADOS E MIGRAÇÕES (SQLite)
# -----------------------------------------------------------------------------
def get_connection():
    conn = sqlite3.connect('orcamentos.db', check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def hash_senha(senha_raw):
    return hashlib.sha256(senha_raw.encode('utf-8')).hexdigest()

def init_db():
    conn = get_connection()
    c = conn.cursor()
    
    # Configurações da Empresa
    c.execute("""
        CREATE TABLE IF NOT EXISTS config_empresa (
            id INTEGER PRIMARY KEY,
            nome_empresa TEXT,
            cnpj TEXT,
            ie TEXT,
            endereco TEXT,
            telefone TEXT,
            email TEXT,
            logo_path TEXT,
            logo_largura REAL DEFAULT 5.0,
            senha_hash TEXT
        )
    """)
    
    c.execute("SELECT COUNT(*) FROM config_empresa")
    if c.fetchone()[0] == 0:
        c.execute("""
            INSERT INTO config_empresa (id, nome_empresa, cnpj, ie, endereco, telefone, email, logo_path, logo_largura, senha_hash)
            VALUES (1, 'Fábrica de Móveis Laurenti Ltda', '44.331.015/0001-08', '186000158114', 
                    'Rua Henrique Villa, 59- Jardim Maria Emília. CEP: 15960000, ARIRANHA-SP', 
                    '(17) 3576-1464', 'contato@laurentimoveis.com.br', '', 5.0, ?)
        """, (hash_senha("laurenti2026"),))

    # Migração logo_largura e senha_hash
    c.execute("PRAGMA table_info(config_empresa)")
    colunas_config = [col[1] for col in c.fetchall()]
    if 'logo_largura' not in colunas_config:
        c.execute("ALTER TABLE config_empresa ADD COLUMN logo_largura REAL DEFAULT 5.0")
    if 'senha_hash' not in colunas_config:
        c.execute("ALTER TABLE config_empresa ADD COLUMN senha_hash TEXT")
        c.execute("UPDATE config_empresa SET senha_hash = ? WHERE id = 1", (hash_senha("laurenti2026"),))

    # Status Comercial Dinâmico
    c.execute("""
        CREATE TABLE IF NOT EXISTS status_comercial (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT UNIQUE NOT NULL
        )
    """)
    
    c.execute("SELECT COUNT(*) FROM status_comercial")
    if c.fetchone()[0] == 0:
        status_iniciais = [("Em Análise",), ("Aprovado",), ("Em Produção",), ("Perdido",)]
        c.executemany("INSERT INTO status_comercial (nome) VALUES (?)", status_iniciais)

    # Consultores
    c.execute("""
        CREATE TABLE IF NOT EXISTS consultores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL
        )
    """)
    
    c.execute("SELECT COUNT(*) FROM consultores")
    if c.fetchone()[0] == 0:
        c.executemany("INSERT INTO consultores (nome) VALUES (?)", [("KATIA LUCIA LOURENCO",), ("Sem Consultor",)])

    # Orçamentos
    c.execute("""
        CREATE TABLE IF NOT EXISTS orcamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            proposta_num INTEGER,
            cliente TEXT NOT NULL,
            contato TEXT,
            tipo_contato TEXT,
            telefone TEXT,
            email TEXT,
            consultor TEXT,
            data TEXT NOT NULL,
            dias_validade INTEGER NOT NULL,
            validade TEXT NOT NULL,
            prazo_entrega TEXT,
            condicoes_pagamento TEXT,
            observacoes TEXT,
            total_liquido REAL DEFAULT 0,
            total_com_opcionais REAL DEFAULT 0,
            status TEXT DEFAULT 'Em Análise'
        )
    """)
    
    c.execute("PRAGMA table_info(orcamentos)")
    colunas_orc = [col[1] for col in c.fetchall()]
    if 'status' not in colunas_orc:
        c.execute("ALTER TABLE orcamentos ADD COLUMN status TEXT DEFAULT 'Em Análise'")

    # Ambientes
    c.execute("""
        CREATE TABLE IF NOT EXISTS ambientes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            orcamento_id INTEGER,
            ordem INTEGER,
            nome_ambiente TEXT NOT NULL,
            especificacoes TEXT,
            total_ambiente REAL DEFAULT 0,
            FOREIGN KEY(orcamento_id) REFERENCES orcamentos(id) ON DELETE CASCADE
        )
    """)

    # Itens / Subitens
    c.execute("""
        CREATE TABLE IF NOT EXISTS itens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ambiente_id INTEGER,
            ordem INTEGER,
            descricao TEXT NOT NULL,
            valor REAL DEFAULT 0,
            eh_opcional INTEGER DEFAULT 0,
            FOREIGN KEY(ambiente_id) REFERENCES ambientes(id) ON DELETE CASCADE
        )
    """)
    
    conn.commit()

init_db()

# Funções Auxiliares protegidas
def get_config():
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM config_empresa WHERE id = 1")
    row = c.fetchone()
    if row:
        d = dict(row)
        d.setdefault('nome_empresa', 'Fábrica de Móveis Laurenti Ltda')
        d.setdefault('cnpj', '44.331.015/0001-08')
        d.setdefault('ie', '186000158114')
        d.setdefault('endereco', 'Rua Henrique Villa, 59- Jardim Maria Emília. CEP: 15960000, ARIRANHA-SP')
        d.setdefault('telefone', '(17) 3576-1464')
        d.setdefault('email', 'contato@laurentimoveis.com.br')
        d.setdefault('logo_path', '')
        d.setdefault('logo_largura', 5.0)
        d.setdefault('senha_hash', hash_senha("laurenti2026"))
        return d
    return {'nome_empresa': 'Fábrica de Móveis Laurenti Ltda', 'cnpj': '44.331.015/0001-08', 'ie': '186000158114', 'endereco': '', 'telefone': '', 'email': '', 'logo_path': '', 'logo_largura': 5.0, 'senha_hash': hash_senha("laurenti2026")}

def get_status_list():
    conn = get_connection()
    df = pd.read_sql_query("SELECT * FROM status_comercial ORDER BY id ASC", conn)
    if not df.empty:
        return df['nome'].tolist()
    return ["Em Análise"]

def get_consultores():
    conn = get_connection()
    return pd.read_sql_query("SELECT * FROM consultores ORDER BY nome ASC", conn)

def get_proxima_proposta():
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT proposta_num FROM orcamentos ORDER BY id DESC LIMIT 1")
    row = c.fetchone()
    if row and row['proposta_num'] is not None:
        try:
            val = int(row['proposta_num'])
            return val + 1
        except ValueError:
            return 1
    return 1

def get_ultimas_condicoes():
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT consultor, prazo_entrega, condicoes_pagamento, observacoes, status FROM orcamentos ORDER BY id DESC LIMIT 1")
    row = c.fetchone()
    status_disponiveis = get_status_list()
    default_status = status_disponiveis[0] if status_disponiveis else "Em Análise"
    
    if row:
        st_val = row['status'] if row['status'] in status_disponiveis else default_status
        return {
            'consultor': row['consultor'] or 'Sem Consultor', 
            'prazo_entrega': row['prazo_entrega'] or '120 dias após medições finais.', 
            'condicoes_pagamento': row['condicoes_pagamento'] or '8 PARCELAS', 
            'observacoes': row['observacoes'] or 'Especificações Gerais: MDF externo cores a definir 18mm / MDF interno Branco TX 18mm',
            'status': st_val
        }
    return {
        'consultor': 'Sem Consultor', 
        'prazo_entrega': '120 dias após medições finais.', 
        'condicoes_pagamento': '8 PARCELAS', 
        'observacoes': 'Especificações Gerais: MDF externo cores a definir 18mm / MDF interno Branco TX 18mm',
        'status': default_status
    }

# Canvas Dinâmico para PDF
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        config = get_config()
        self.saveState()
        self.setFont("Helvetica", 7)
        self.setFillColor(colors.HexColor("#333333"))
        
        self.setStrokeColor(colors.HexColor("#BDC3C7"))
        self.setLineWidth(0.5)
        self.line(1.2*cm, 1.8*cm, A4[0] - 1.2*cm, 1.8*cm)
        
        line1 = f"{config.get('nome_empresa', '')} CNPJ: {config.get('cnpj', '')} IE: {config.get('ie', '')}"
        line2 = f"{config.get('endereco', '')} / Fone: {config.get('telefone', '')}"
        
        self.drawString(1.2*cm, 1.3*cm, line1)
        self.drawString(1.2*cm, 0.9*cm, line2)
        
        data_hora_str = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        page_str = f"Página {self._pageNumber} de {page_count}"
        
        self.drawRightString(A4[0] - 1.2*cm, 1.3*cm, page_str)
        self.drawRightString(A4[0] - 1.2*cm, 0.9*cm, data_hora_str)
        self.restoreState()

# -----------------------------------------------------------------------------
# 2. GERADORES DE ARQUIVOS (PDF, WORD, EXCEL)
# -----------------------------------------------------------------------------
def gerar_pdf_orcamento(cliente_info, ambientes_list):
    config = get_config()
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.2*cm,
        leftMargin=1.2*cm,
        topMargin=1.2*cm,
        bottomMargin=2.2*cm
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontSize=8.5, leading=13, textColor=colors.HexColor('#222222'))
    body_bold = ParagraphStyle('BodyBold', parent=body_style, fontName='Helvetica-Bold')
    right_bold = ParagraphStyle('RightBold', parent=body_bold, alignment=2)
    header_title = ParagraphStyle('HeaderTitle', parent=body_bold, fontSize=9, textColor=colors.HexColor('#1E293B'))
    amb_title_style = ParagraphStyle('AmbTitleStyle', parent=body_bold, fontSize=9.5, textColor=colors.HexColor('#0F172A'))

    logo_largura_cm = float(config.get('logo_largura', 5.0))
    col_logo = Paragraph("<b>LAURENTI MÓVEIS</b>", body_bold)
    logo_p = config.get('logo_path', '')
    
    if logo_p and os.path.exists(logo_p):
        try:
            with Image.open(logo_p) as img:
                w_orig, h_orig = img.size
                max_w = logo_largura_cm * cm
                max_h = 3.0 * cm
                ratio = min(max_w / w_orig, max_h / h_orig)
                final_w = w_orig * ratio
                final_h = h_orig * ratio
                col_logo = RLImage(logo_p, width=final_w, height=final_h)
        except Exception:
            pass

    cli_text = f"""
    <b>Cliente:</b> {cliente_info.get('cliente', '')}<br/>
    <b>Contato:</b> {cliente_info.get('contato', '')}<br/>
    <b>Consultor:</b> {cliente_info.get('consultor', 'N/A')}<br/>
    <b>Tipo de Contato:</b> {cliente_info.get('tipo_contato', 'Residencial')}<br/>
    <b>Telefone:</b> {cliente_info.get('telefone', '')}<br/>
    <b>E-mail:</b> {cliente_info.get('email', '')}
    """
    
    p_num = cliente_info.get('proposta_num', 1)
    prop_num_str = f"{int(p_num):04d}" if str(p_num).isdigit() else str(p_num)
    
    prop_text = f"""
    <b>Proposta:</b> {prop_num_str}<br/><br/>
    <b>Data:</b> {cliente_info.get('data', '')}<br/><br/>
    <b>Validade:</b> {cliente_info.get('validade', '')}
    """
    
    width_col_logo = max(logo_largura_cm + 0.5, 4.0) * cm
    width_col_cli = (14.2 * cm) - width_col_logo
    
    t_head = Table([
        [col_logo, Paragraph(cli_text, body_style), Paragraph(prop_text, body_style)]
    ], colWidths=[width_col_logo, width_col_cli, 4.4*cm])
    
    t_head.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 6),
        ('BACKGROUND', (2,0), (2,0), colors.HexColor('#F8FAFC')),
    ]))
    story.append(t_head)
    story.append(Spacer(1, 10))

    table_data = [
        [Paragraph("<b>Item</b>", header_title), Paragraph("<b>Valor (R$)</b>", ParagraphStyle('HRight', parent=header_title, alignment=2))]
    ]
    
    tot_liquido = 0.0
    tot_opcionais = 0.0
    row_styles = []

    for idx_a, amb in enumerate(ambientes_list):
        ordem_amb = idx_a + 1
        nome_amb_upper = amb['nome'].upper()
        
        itens_normais = [i for i in amb['itens'] if not i['eh_opcional']]
        itens_opcionais = [i for i in amb['itens'] if i['eh_opcional']]
        
        tot_amb = sum(float(i['valor']) for i in itens_normais)
        tot_liquido += tot_amb
        tot_amb_str = f"R$ {tot_amb:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        current_row = len(table_data)
        amb_header_text = f"<b>{ordem_amb}- {nome_amb_upper}</b>"
        table_data.append([Paragraph(amb_header_text, amb_title_style), Paragraph(f"<b>{tot_amb_str}</b>", right_bold)])
        row_styles.append(('BACKGROUND', (0, current_row), (-1, current_row), colors.HexColor('#F1F5F9')))
        
        espec_fmt = amb.get('especificacoes', '').replace('\n', '<br/>')
        desc_amb_text = f"<i>{espec_fmt}</i>" if espec_fmt else ""
        
        for item in itens_normais:
            val_f = f"{float(item['valor']):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            desc_item = item['descricao'].replace('\n', '<br/>&nbsp;&nbsp;&nbsp;&nbsp;')
            desc_amb_text += f"<br/>- {desc_item} - R$ {val_f}"
            
        if itens_opcionais:
            desc_amb_text += "<br/><br/><b>• Acessórios opcionais a serem acrescidos:</b>"
            for item in itens_opcionais:
                val_opc = float(item['valor'])
                tot_opcionais += val_opc
                val_f = f"{val_opc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                desc_item = item['descricao'].replace('\n', '<br/>&nbsp;&nbsp;&nbsp;&nbsp;')
                desc_amb_text += f"<br/>- {desc_item}, acréscimo - R$ {val_f}"

        table_data.append([Paragraph(desc_amb_text, body_style), Paragraph("", body_style)])

    t_itens = Table(table_data, colWidths=[14.6*cm, 4*cm])
    
    t_style = [
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#E2E8F0')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#94A3B8')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 7),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ] + row_styles
    
    t_itens.setStyle(TableStyle(t_style))
    story.append(t_itens)
    story.append(Spacer(1, 10))

    tot_com_opc = tot_liquido + tot_opcionais
    tot_liquido_str = f"R$ {tot_liquido:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    tot_com_opc_str = f"R$ {tot_com_opc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    totais_box_text = f"""
    <font size=9 color='#64748B'><b>Total Líquido:</b></font> &nbsp;&nbsp;&nbsp;&nbsp; <font size=11 color='#0F172A'><b>{tot_liquido_str}</b></font><br/>
    <font size=8 color='#64748B'>Total c/ Opcionais:</font> &nbsp;&nbsp;&nbsp;&nbsp; <font size=10 color='#D97706'><b>{tot_com_opc_str}</b></font>
    """
    
    t_totais = Table([
        ["", Paragraph(totais_box_text, right_bold)]
    ], colWidths=[10.6*cm, 8.0*cm])
    
    t_totais.setStyle(TableStyle([
        ('BOX', (1,0), (1,0), 0.5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (1,0), (1,0), colors.HexColor('#F1F5F9')),
        ('PADDING', (1,0), (1,0), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t_totais)
    story.append(Spacer(1, 8))

    obs_fmt = cliente_info.get('observacoes', '').replace('\n', '<br/>')
    cond_text = f"""
    <b>Prazo de Entrega:</b> {cliente_info.get('prazo_entrega', '')}<br/>
    <b>Condição de Pagamento:</b> {cliente_info.get('condicoes_pagamento', '')}<br/>
    <b>Observações:</b> {obs_fmt}
    """
    
    t_cond = Table([
        [Paragraph(cond_text, body_style)]
    ], colWidths=[18.6*cm])
    
    t_cond.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(t_cond)

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def gerar_word_orcamento(cliente_info, ambientes_list):
    config = get_config()
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.47)
        section.bottom_margin = Inches(0.47)
        section.left_margin = Inches(0.47)
        section.right_margin = Inches(0.47)

    # 1. CABEÇALHO
    t_head = doc.add_table(rows=1, cols=3)
    t_head.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_head.autofit = False
    set_table_borders(t_head, color="CBD5E1")
    
    logo_largura_cm = float(config.get('logo_largura', 5.0))
    w_logo = max(logo_largura_cm / 2.54, 1.8)
    
    widths = [Inches(w_logo), Inches(5.5 - w_logo), Inches(1.8)]
    for row in t_head.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            set_cell_margins(cell, top=120, bottom=120, left=120, right=120)

    cell_logo, cell_cli, cell_prop = t_head.rows[0].cells
    set_cell_background(cell_prop, "F8FAFC")

    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p = config.get('logo_path', '')
    if logo_p and os.path.exists(logo_p):
        try:
            p_logo.add_run().add_picture(logo_p, width=Inches(logo_largura_cm / 2.54))
        except Exception:
            p_logo.add_run("LAURENTI MÓVEIS").bold = True
    else:
        p_logo.add_run("LAURENTI MÓVEIS").bold = True

    p_cli = cell_cli.paragraphs[0]
    p_cli.paragraph_format.line_spacing = 1.15
    
    def add_line(p, label, value):
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(f"{value}\n")

    add_line(p_cli, "Cliente", cliente_info.get('cliente', ''))
    add_line(p_cli, "Contato", cliente_info.get('contato', ''))
    add_line(p_cli, "Consultor", cliente_info.get('consultor', ''))
    add_line(p_cli, "Tipo de Contato", cliente_info.get('tipo_contato', ''))
    add_line(p_cli, "Telefone", cliente_info.get('telefone', ''))
    
    r_em_lbl = p_cli.add_run("E-mail: ")
    r_em_lbl.bold = True
    p_cli.add_run(f"{cliente_info.get('email', '')}")

    p_prop = cell_prop.paragraphs[0]
    p_num = cliente_info.get('proposta_num', 1)
    prop_str = f"{int(p_num):04d}" if str(p_num).isdigit() else str(p_num)
    
    add_line(p_prop, "Proposta", prop_str)
    p_prop.add_run("\n")
    add_line(p_prop, "Data", cliente_info.get('data', ''))
    p_prop.add_run("\n")
    
    r_val_lbl = p_prop.add_run("Validade: ")
    r_val_lbl.bold = True
    p_prop.add_run(f"{cliente_info.get('validade', '')}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. TABELA DE ITENS
    t_itens = doc.add_table(rows=1, cols=2)
    t_itens.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_itens.autofit = False
    set_table_borders(t_itens, color="CBD5E1")

    t_itens.columns[0].width = Inches(5.7)
    t_itens.columns[1].width = Inches(1.6)

    hdr_cells = t_itens.rows[0].cells
    set_cell_background(hdr_cells[0], "E2E8F0")
    set_cell_background(hdr_cells[1], "E2E8F0")
    set_cell_margins(hdr_cells[0], top=100, bottom=100, left=120, right=120)
    set_cell_margins(hdr_cells[1], top=100, bottom=100, left=120, right=120)

    hdr_cells[0].paragraphs[0].add_run("Item").bold = True
    p_h_val = hdr_cells[1].paragraphs[0]
    p_h_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_h_val.add_run("Valor (R$)").bold = True

    tot_liquido = 0.0
    tot_opcionais = 0.0

    for idx_a, amb in enumerate(ambientes_list):
        ordem_amb = idx_a + 1
        itens_normais = [i for i in amb['itens'] if not i['eh_opcional']]
        itens_opcionais = [i for i in amb['itens'] if i['eh_opcional']]
        
        tot_amb = sum(float(i['valor']) for i in itens_normais)
        tot_liquido += tot_amb
        tot_amb_str = f"R$ {tot_amb:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        row_amb = t_itens.add_row()
        c_tit, c_val = row_amb.cells
        set_cell_background(c_tit, "F1F5F9")
        set_cell_background(c_val, "F1F5F9")
        set_cell_margins(c_tit, top=100, bottom=100, left=120, right=120)
        set_cell_margins(c_val, top=100, bottom=100, left=120, right=120)

        p_tit = c_tit.paragraphs[0]
        r_tit = p_tit.add_run(f"{ordem_amb}- {amb['nome'].upper()}")
        r_tit.bold = True
        
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_val = p_val.add_run(tot_amb_str)
        r_val.bold = True

        row_desc = t_itens.add_row()
        c_desc, c_empty = row_desc.cells
        set_cell_margins(c_desc, top=100, bottom=100, left=120, right=120)

        p_desc = c_desc.paragraphs[0]
        p_desc.paragraph_format.line_spacing = 1.15
        
        if amb.get('especificacoes'):
            r_esp = p_desc.add_run(f"{amb['especificacoes']}\n")
            r_esp.italic = True

        for item in itens_normais:
            val_f = f"{float(item['valor']):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            p_desc.add_run(f"- {item['descricao']} - R$ {val_f}\n")

        if itens_opcionais:
            p_desc.add_run("\n• Acessórios opcionais a serem acrescidos:\n").bold = True
            for item in itens_opcionais:
                val_opc = float(item['valor'])
                tot_opcionais += val_opc
                val_f = f"{val_opc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                p_desc.add_run(f"- {item['descricao']}, acréscimo - R$ {val_f}\n")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 3. TOTAIS
    tot_com_opc = tot_liquido + tot_opcionais
    tot_liq_str = f"R$ {tot_liquido:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    tot_opc_str = f"R$ {tot_com_opc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    t_tot = doc.add_table(rows=1, cols=2)
    t_tot.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tot.autofit = False
    
    t_tot.columns[0].width = Inches(4.3)
    t_tot.columns[1].width = Inches(3.0)

    cell_blank, cell_box = t_tot.rows[0].cells
    set_cell_background(cell_box, "F1F5F9")
    set_cell_margins(cell_box, top=120, bottom=120, left=120, right=120)
    
    p_box = cell_box.paragraphs[0]
    p_box.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    r_tl1 = p_box.add_run("Total Líquido:  ")
    r_tl1.font.size = Pt(9)
    r_tl1.bold = True
    r_tl2 = p_box.add_run(f"{tot_liq_str}\n")
    r_tl2.font.size = Pt(11)
    r_tl2.bold = True
    
    r_to1 = p_box.add_run("Total c/ Opcionais:  ")
    r_to1.font.size = Pt(8)
    r_to2 = p_box.add_run(f"{tot_opc_str}")
    r_to2.font.size = Pt(10)
    r_to2.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 4. OBSERVAÇÕES
    t_cond = doc.add_table(rows=1, cols=1)
    t_cond.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_cond.autofit = False
    t_cond.columns[0].width = Inches(7.3)
    
    set_table_borders(t_cond, color="CBD5E1")
    cell_cond = t_cond.rows[0].cells[0]
    set_cell_background(cell_cond, "F8FAFC")
    set_cell_margins(cell_cond, top=120, bottom=120, left=120, right=120)

    p_cond = cell_cond.paragraphs[0]
    p_cond.paragraph_format.line_spacing = 1.15
    
    add_line(p_cond, "Prazo de Entrega", cliente_info.get('prazo_entrega', ''))
    add_line(p_cond, "Condição de Pagamento", cliente_info.get('condicoes_pagamento', ''))
    
    r_obs_lbl = p_cond.add_run("Observações: ")
    r_obs_lbl.bold = True
    p_cond.add_run(f"{cliente_info.get('observacoes', '')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_excel_orcamento(cliente_info, ambientes_list):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Orçamento"
    
    ws.views.sheetView[0].showGridLines = True

    font_header = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    font_bold = Font(name="Calibri", size=11, bold=True)
    font_title = Font(name="Calibri", size=14, bold=True, color="1E293B")
    fill_header = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    fill_amb = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
    fill_tot = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    align_right = Alignment(horizontal="right", vertical="top")
    align_left = Alignment(horizontal="left", vertical="top", wrap_text=True)

    config = get_config()
    p_num = cliente_info.get('proposta_num', 1)
    prop_str = f"{int(p_num):04d}" if str(p_num).isdigit() else str(p_num)

    ws["A1"] = config.get('nome_empresa', 'LAURENTI MÓVEIS')
    ws["A1"].font = font_title
    
    ws["A3"] = f"Cliente: {cliente_info.get('cliente', '')}"
    ws["A3"].font = font_bold
    ws["A4"] = f"Contato: {cliente_info.get('contato', '')}"
    ws["A5"] = f"Telefone: {cliente_info.get('telefone', '')}"
    ws["A6"] = f"E-mail: {cliente_info.get('email', '')}"

    ws["C3"] = f"Proposta Nº: #{prop_str}"
    ws["C3"].font = font_bold
    ws["C4"] = f"Data: {cliente_info.get('data', '')}"
    ws["C5"] = f"Validade: {cliente_info.get('validade', '')}"
    ws["C6"] = f"Consultor: {cliente_info.get('consultor', '')}"

    row_idx = 8
    ws.cell(row=row_idx, column=1, value="Ambiente / Subitem").font = font_header
    ws.cell(row=row_idx, column=1).fill = fill_header
    ws.cell(row=row_idx, column=2, value="Tipo").font = font_header
    ws.cell(row=row_idx, column=2).fill = fill_header
    ws.cell(row=row_idx, column=3, value="Valor (R$)").font = font_header
    ws.cell(row=row_idx, column=3).fill = fill_header
    ws.cell(row=row_idx, column=3).alignment = align_right

    row_idx += 1
    tot_liquido = 0.0
    tot_opcionais = 0.0

    for idx_a, amb in enumerate(ambientes_list):
        ordem_amb = idx_a + 1
        itens_normais = [i for i in amb['itens'] if not i['eh_opcional']]
        itens_opcionais = [i for i in amb['itens'] if i['eh_opcional']]
        tot_amb = sum(float(i['valor']) for i in itens_normais)
        tot_liquido += tot_amb

        ws.cell(row=row_idx, column=1, value=f"{ordem_amb}- {amb['nome'].upper()}").font = font_bold
        ws.cell(row=row_idx, column=1).fill = fill_amb
        ws.cell(row=row_idx, column=2, value="Ambiente").font = font_bold
        ws.cell(row=row_idx, column=2).fill = fill_amb
        
        c_val = ws.cell(row=row_idx, column=3, value=tot_amb)
        c_val.font = font_bold
        c_val.number_format = 'R$ #,##0.00'
        c_val.fill = fill_amb
        c_val.alignment = align_right
        row_idx += 1

        if amb.get('especificacoes'):
            ws.cell(row=row_idx, column=1, value=f"Especificações: {amb['especificacoes']}").alignment = align_left
            row_idx += 1

        for item in itens_normais:
            ws.cell(row=row_idx, column=1, value=f"- {item['descricao']}").alignment = align_left
            ws.cell(row=row_idx, column=2, value="Normal")
            c_v = ws.cell(row=row_idx, column=3, value=float(item['valor']))
            c_v.number_format = 'R$ #,##0.00'
            c_v.alignment = align_right
            row_idx += 1

        for item in itens_opcionais:
            tot_opcionais += float(item['valor'])
            ws.cell(row=row_idx, column=1, value=f"- [OPCIONAL] {item['descricao']}").alignment = align_left
            ws.cell(row=row_idx, column=2, value="Opcional")
            c_v = ws.cell(row=row_idx, column=3, value=float(item['valor']))
            c_v.number_format = 'R$ #,##0.00'
            c_v.alignment = align_right
            row_idx += 1
        
        row_idx += 1

    ws.cell(row=row_idx, column=2, value="Total Líquido:").font = font_bold
    c_tl = ws.cell(row=row_idx, column=3, value=tot_liquido)
    c_tl.font = font_bold
    c_tl.number_format = 'R$ #,##0.00'
    c_tl.alignment = align_right
    row_idx += 1

    ws.cell(row=row_idx, column=2, value="Total c/ Opcionais:").font = font_bold
    c_to = ws.cell(row=row_idx, column=3, value=(tot_liquido + tot_opcionais))
    c_to.font = font_bold
    c_to.number_format = 'R$ #,##0.00'
    c_to.alignment = align_right
    row_idx += 2

    ws.cell(row=row_idx, column=1, value=f"Prazo de Entrega: {cliente_info.get('prazo_entrega', '')}").font = font_bold
    row_idx += 1
    ws.cell(row=row_idx, column=1, value=f"Condição de Pagamento: {cliente_info.get('condicoes_pagamento', '')}").font = font_bold
    row_idx += 1
    ws.cell(row=row_idx, column=1, value=f"Observações: {cliente_info.get('observacoes', '')}")

    ws.column_dimensions['A'].width = 65
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# 3. CONTROLE DE AUTENTICAÇÃO E SESSÃO
# -----------------------------------------------------------------------------
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

def tela_login():
    st.markdown("<h2 style='text-align: center;'>🔐 Acesso ao Sistema Laurenti Móveis</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #64748B;'>Digite a senha de segurança para continuar</p>", unsafe_allow_html=True)
    
    col_l1, col_l2, col_l3 = st.columns([1, 1.2, 1])
    with col_l2:
        with st.form("form_login"):
            senha_input = st.text_input("Senha", type="password", placeholder="Digite sua senha")
            btn_entrar = st.form_submit_button("🔓 Entrar no Sistema", use_container_width=True)
            
            if btn_entrar:
                config = get_config()
                if hash_senha(senha_input) == config.get('senha_hash'):
                    st.session_state.autenticado = True
                    st.success("Acesso liberado!")
                    st.rerun()
                else:
                    st.error("Senha incorreta. Tente novamente.")

if not st.session_state.autenticado:
    tela_login()
    st.stop()

# -----------------------------------------------------------------------------
# 4. INTERFACE PRINCIPAL STREAMLIT
# -----------------------------------------------------------------------------
if 'ambientes' not in st.session_state:
    st.session_state.ambientes = []

if 'edit_index' not in st.session_state:
    st.session_state.edit_index = None

if 'confirm_del' not in st.session_state:
    st.session_state.confirm_del = None

if 'form_version' not in st.session_state:
    st.session_state.form_version = 1

if 'expand_ambientes' not in st.session_state:
    st.session_state.expand_ambientes = True

def reseta_formulario_limpo():
    ultimas_cond = get_ultimas_condicoes()
    st.session_state.ambientes = []
    st.session_state.edit_index = None
    
    st.session_state.cli_nome = ""
    st.session_state.cli_contato = ""
    st.session_state.cli_tel = ""
    st.session_state.cli_email = ""
    st.session_state.cli_tipo = "Residencial"
    st.session_state.cli_consultor = ultimas_cond['consultor']
    st.session_state.cli_prazo = ultimas_cond['prazo_entrega']
    st.session_state.cli_cond = ultimas_cond['condicoes_pagamento']
    st.session_state.cli_obs = ultimas_cond['observacoes']
    st.session_state.cli_prop = get_proxima_proposta()
    st.session_state.cli_status = ultimas_cond['status']
    
    st.session_state.expand_ambientes = True
    st.session_state.form_version += 1

st.sidebar.markdown("### 👤 Usuário Autenticado")
if st.sidebar.button("🚪 Sair / Bloquear Sistema", use_container_width=True):
    st.session_state.autenticado = False
    st.rerun()

st.sidebar.markdown("---")

opcoes_menu = ["➕ Novo / Editar Orçamento", "📋 Orçamentos Salvos", "⚙️ Configurações"]

if 'radio_menu' not in st.session_state:
    st.session_state.radio_menu = opcoes_menu[0]

if 'change_tab_to' in st.session_state and st.session_state.change_tab_to:
    st.session_state.radio_menu = st.session_state.change_tab_to
    st.session_state.change_tab_to = None

menu = st.sidebar.radio("Navegação", opcoes_menu, key="radio_menu")

def recalcular_totais():
    total_liquido = 0.0
    total_opcionais = 0.0
    for amb in st.session_state.ambientes:
        tot_amb = 0.0
        for item in amb['itens']:
            val = float(item['valor'])
            if not item['eh_opcional']:
                tot_amb += val
            else:
                total_opcionais += val
        amb['total_ambiente'] = tot_amb
        total_liquido += tot_amb
    return total_liquido, (total_liquido + total_opcionais)

st.title("🏭 Laurenti Móveis — Gestão de Orçamentos")

# --- ABA 1: NOVO / EDITAR ORÇAMENTO ---
if menu == "➕ Novo / Editar Orçamento":
    st.subheader("Formulário de Orçamento")
    
    ultimas_cond = get_ultimas_condicoes()
    df_cons = get_consultores()
    consultores_opts = df_cons['nome'].tolist() if not df_cons.empty else ["Sem Consultor"]
    status_opts = get_status_list()

    if st.button("✨ Criar Novo Orçamento Limpo"):
        reseta_formulario_limpo()
        st.rerun()

    prop_num_atual = st.session_state.get('cli_prop', get_proxima_proposta())
    prop_formatted = f"{int(prop_num_atual):04d}" if str(prop_num_atual).isdigit() else str(prop_num_atual)
    
    v = st.session_state.form_version

    with st.expander("👤 Dados do Cliente, Status e Proposta", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            cliente = st.text_input("Cliente *", value=st.session_state.get('cli_nome', ''), placeholder="Nome da pessoa ou empresa", key=f"in_cli_nome_{v}")
            st.session_state.cli_nome = cliente
            
            contato = st.text_input("Contato", value=st.session_state.get('cli_contato', ''), placeholder="Nome do responsável", key=f"in_cli_contato_{v}")
            st.session_state.cli_contato = contato
            
            tipo_opts = ["Residencial", "Comercial", "Arquitetura", "Outros"]
            idx_tipo = tipo_opts.index(st.session_state.get('cli_tipo', 'Residencial')) if st.session_state.get('cli_tipo') in tipo_opts else 0
            tipo_contato = st.selectbox("Tipo de Contato", tipo_opts, index=idx_tipo, key=f"in_cli_tipo_{v}")
            st.session_state.cli_tipo = tipo_contato
            
            cons_val = st.session_state.get('cli_consultor', ultimas_cond['consultor'])
            idx_cons_padrao = consultores_opts.index(cons_val) if cons_val in consultores_opts else 0
            consultor = st.selectbox("Consultor *", consultores_opts, index=idx_cons_padrao, key=f"in_cli_cons_{v}")
            st.session_state.cli_consultor = consultor
        
        with col2:
            telefone = st.text_input("Telefone", value=st.session_state.get('cli_tel', ''), placeholder="(00) 00000-0000", key=f"in_cli_tel_{v}")
            st.session_state.cli_tel = telefone
            
            email = st.text_input("E-mail", value=st.session_state.get('cli_email', ''), placeholder="cliente@email.com", key=f"in_cli_email_{v}")
            st.session_state.cli_email = email
            
            data_atual = st.date_input("Data da Proposta", value=datetime.now().date(), key=f"in_cli_data_{v}")
            
            status_val = st.session_state.get('cli_status', ultimas_cond['status'])
            idx_status = status_opts.index(status_val) if status_val in status_opts else 0
            status_orcamento = st.selectbox("Status Comercial", status_opts, index=idx_status, key=f"in_cli_status_{v}")
            st.session_state.cli_status = status_orcamento
            
        with col3:
            st.text_input("Proposta Nº (Automático)", value=prop_formatted, disabled=True)
            dias_validade = st.radio("Validade em Dias", [7, 10, 15, 30], index=3, horizontal=True, key=f"in_cli_val_{v}")
            data_validade = data_atual + timedelta(days=dias_validade)
            st.info(f"📅 **Validade:** {data_validade.strftime('%d/%m/%Y')}")
            
        col_cond1, col_cond2 = st.columns(2)
        with col_cond1:
            prazo_entrega = st.text_input("Prazo de Entrega", value=st.session_state.get('cli_prazo', ultimas_cond['prazo_entrega']), key=f"in_cli_prazo_{v}")
            st.session_state.cli_prazo = prazo_entrega
        with col_cond2:
            condicoes_pagamento = st.text_input("Condições de Pagamento", value=st.session_state.get('cli_cond', ultimas_cond['condicoes_pagamento']), key=f"in_cli_cond_{v}")
            st.session_state.cli_cond = condicoes_pagamento
            
        observacoes = st.text_area("Observações Gerais", value=st.session_state.get('cli_obs', ultimas_cond['observacoes']), key=f"in_cli_obs_{v}")
        st.session_state.cli_obs = observacoes

    st.markdown("---")
    
    col_amb_head1, col_amb_head2 = st.columns([3, 1])
    col_amb_head1.subheader("🛋️ Ambientes e Subitens")
    
    btn_toggle_text = "📁 Contrair Todos os Ambientes" if st.session_state.expand_ambientes else "📂 Expandir Todos os Ambientes"
    if col_amb_head2.button(btn_toggle_text, use_container_width=True):
        st.session_state.expand_ambientes = not st.session_state.expand_ambientes
        st.rerun()

    with st.form("form_novo_ambiente", clear_on_submit=True):
        col_amb1, col_amb2 = st.columns([2, 3])
        with col_amb1:
            nome_amb = st.text_input("Nome do Ambiente *", placeholder="Ex: COZINHA, GOURMET, SALA")
        with col_amb2:
            espec_amb = st.text_area("Especificações do Ambiente", placeholder="Ex: MDF Nogueira Ambar / MDF Nude Vel\n(Pressione Enter para quebrar linha)", height=68)
        
        btn_add_amb = st.form_submit_button("➕ Adicionar Ambiente")
        if btn_add_amb:
            if nome_amb:
                st.session_state.ambientes.append({
                    'nome': nome_amb,
                    'especificacoes': espec_amb,
                    'total_ambiente': 0.0,
                    'itens': []
                })
                st.session_state.expand_ambientes = True
                st.rerun()

    if st.session_state.ambientes:
        for idx_amb, amb in enumerate(st.session_state.ambientes):
            ordem_amb = idx_amb + 1
            
            for item in amb['itens']:
                if 'id' not in item:
                    item['id'] = str(uuid.uuid4())

            amb['itens'] = sorted(amb['itens'], key=lambda x: x['eh_opcional'])

            with st.expander(f"🛋️ **Item {ordem_amb}: {amb['nome'].upper()}** — Total: R$ {amb['total_ambiente']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."), expanded=st.session_state.expand_ambientes):
                
                col_e1, col_e2, col_e3 = st.columns([2, 3, 1])
                with col_e1:
                    amb['nome'] = st.text_input(f"Nome #{ordem_amb}", value=amb['nome'], key=f"edit_nome_{v}_{idx_amb}")
                with col_e2:
                    amb['especificacoes'] = st.text_area(f"Especificações #{ordem_amb}", value=amb['especificacoes'], key=f"edit_espec_{v}_{idx_amb}", height=68)
                with col_e3:
                    st.write(" ")
                    if st.button("🗑️ Remover Ambiente", key=f"del_amb_btn_{v}_{idx_amb}"):
                        st.session_state.confirm_del = f"amb_{idx_amb}"

                if st.session_state.confirm_del == f"amb_{idx_amb}":
                    st.warning("⚠️ Confirma a exclusão deste ambiente e todos seus itens?")
                    c_del1, c_del2 = st.columns(2)
                    if c_del1.button("✅ Confirmar Exclusão", key=f"conf_del_amb_{v}_{idx_amb}"):
                        st.session_state.ambientes.pop(idx_amb)
                        st.session_state.confirm_del = None
                        st.rerun()
                    if c_del2.button("❌ Cancelar", key=f"canc_del_amb_{v}_{idx_amb}"):
                        st.session_state.confirm_del = None
                        st.rerun()

                st.markdown(f"##### Subitens do Ambiente {ordem_amb}")
                
                with st.form(f"form_add_subitem_{v}_{idx_amb}", clear_on_submit=True):
                    col_i1, col_i2, col_i3, col_i4 = st.columns([3, 1.5, 1, 1])
                    with col_i1:
                        desc_sub = st.text_area("Descrição do Subitem", placeholder="Ex: Armário aéreo 3,67m\n8 portas basculantes", height=68)
                    with col_i2:
                        val_sub = st.number_input("Valor (R$)", min_value=0.0, step=100.0, format="%.2f")
                    with col_i3:
                        st.write(" ")
                        opc_sub = st.checkbox("Opcional?")
                    with col_i4:
                        st.write(" ")
                        btn_sub = st.form_submit_button("➕ Incluir Subitem")
                        
                    if btn_sub:
                        if desc_sub and val_sub > 0:
                            amb['itens'].append({
                                'id': str(uuid.uuid4()),
                                'descricao': desc_sub,
                                'valor': val_sub,
                                'eh_opcional': opc_sub
                            })
                            st.rerun()

                if amb['itens']:
                    for idx_item, item in enumerate(amb['itens']):
                        item_id = item['id']
                        tag_opc = " (OPCIONAL)" if item['eh_opcional'] else ""
                        
                        col_it1, col_it2, col_it3, col_it4 = st.columns([3, 1.5, 1, 0.5])
                        with col_it1:
                            item['descricao'] = st.text_area(f"Subitem {ordem_amb}.{idx_item+1}{tag_opc}", value=item['descricao'], key=f"item_desc_{item_id}", height=68)
                        with col_it2:
                            item['valor'] = st.number_input("Valor R$", value=float(item['valor']), step=50.0, format="%.2f", key=f"item_val_{item_id}")
                        with col_it3:
                            st.write(" ")
                            item['eh_opcional'] = st.checkbox("Opcional?", value=bool(item['eh_opcional']), key=f"item_opc_{item_id}")
                        with col_it4:
                            st.write(" ")
                            if st.button("❌", key=f"del_item_{item_id}"):
                                amb['itens'].pop(idx_item)
                                st.rerun()

    tot_liquido, tot_com_opcionais = recalcular_totais()
    
    st.markdown("---")
    st.markdown("### 📊 Totais e Emissão de Orçamento")
    c_tot1, c_tot2 = st.columns(2)
    c_tot1.metric("Total Líquido (Sem Opcionais)", f"R$ {tot_liquido:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
    c_tot2.metric("Total Com Opcionais", f"R$ {tot_com_opcionais:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))

    col_btn1, col_btn2 = st.columns(2)
    
    with col_btn1:
        if st.button("💾 Salvar Orçamento no Banco", type="primary", use_container_width=True):
            if not cliente:
                st.error("Por favor, preencha o nome do Cliente.")
            elif not st.session_state.ambientes:
                st.error("Adicione pelo menos um ambiente antes de salvar.")
            else:
                conn = get_connection()
                c = conn.cursor()
                
                num_para_salvar = int(prop_num_atual) if str(prop_num_atual).isdigit() else get_proxima_proposta()

                if st.session_state.edit_index:
                    orc_id = st.session_state.edit_index
                    c.execute("DELETE FROM ambientes WHERE orcamento_id = ?", (orc_id,))
                    c.execute("""
                        UPDATE orcamentos SET proposta_num=?, cliente=?, contato=?, tipo_contato=?, telefone=?, email=?, consultor=?, data=?, dias_validade=?, validade=?, prazo_entrega=?, condicoes_pagamento=?, observacoes=?, total_liquido=?, total_com_opcionais=?, status=?
                        WHERE id=?
                    """, (num_para_salvar, cliente, contato, tipo_contato, telefone, email, consultor, str(data_atual), dias_validade, str(data_validade.strftime('%d/%m/%Y')), prazo_entrega, condicoes_pagamento, observacoes, tot_liquido, tot_com_opcionais, status_orcamento, orc_id))
                else:
                    c.execute("""
                        INSERT INTO orcamentos (proposta_num, cliente, contato, tipo_contato, telefone, email, consultor, data, dias_validade, validade, prazo_entrega, condicoes_pagamento, observacoes, total_liquido, total_com_opcionais, status)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (num_para_salvar, cliente, contato, tipo_contato, telefone, email, consultor, str(data_atual), dias_validade, str(data_validade.strftime('%d/%m/%Y')), prazo_entrega, condicoes_pagamento, observacoes, tot_liquido, tot_com_opcionais, status_orcamento))
                    orc_id = c.lastrowid
                
                for idx_a, amb in enumerate(st.session_state.ambientes):
                    c.execute("""
                        INSERT INTO ambientes (orcamento_id, ordem, nome_ambiente, especificacoes, total_ambiente)
                        VALUES (?, ?, ?, ?, ?)
                    """, (orc_id, idx_a + 1, amb['nome'], amb['especificacoes'], amb['total_ambiente']))
                    
                    amb_id = c.lastrowid
                    for idx_i, item in enumerate(amb['itens']):
                        c.execute("""
                            INSERT INTO itens (ambiente_id, ordem, descricao, valor, eh_opcional)
                            VALUES (?, ?, ?, ?, ?)
                        """, (amb_id, idx_i + 1, item['descricao'], item['valor'], 1 if item['eh_opcional'] else 0))
                
                conn.commit()
                st.success(f"✅ Orçamento Proposta Nº {prop_formatted} salvo com sucesso!")

    with col_btn2:
        cli_info = {
            'proposta_num': prop_formatted,
            'cliente': cliente if cliente else 'CLIENTE NÃO INFORMADO',
            'contato': contato,
            'tipo_contato': tipo_contato,
            'telefone': telefone,
            'email': email,
            'consultor': consultor,
            'data': data_atual.strftime('%d/%m/%Y'),
            'dias_validade': dias_validade,
            'validade': data_validade.strftime('%d/%m/%Y'),
            'prazo_entrega': prazo_entrega,
            'condicoes_pagamento': condicoes_pagamento,
            'observacoes': observacoes
        }
        
        c_exp1, c_exp2, c_exp3 = st.columns(3)
        
        pdf_bytes = gerar_pdf_orcamento(cli_info, st.session_state.ambientes)
        c_exp1.download_button(
            label="📄 PDF",
            data=pdf_bytes,
            file_name=f"Orcamento_{prop_formatted}_{cliente.replace(' ', '_') if cliente else 'Cliente'}.pdf",
            mime="application/pdf",
            use_container_width=True
        )

        word_bytes = gerar_word_orcamento(cli_info, st.session_state.ambientes)
        c_exp2.download_button(
            label="📝 Word",
            data=word_bytes,
            file_name=f"Orcamento_{prop_formatted}_{cliente.replace(' ', '_') if cliente else 'Cliente'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        excel_bytes = gerar_excel_orcamento(cli_info, st.session_state.ambientes)
        c_exp3.download_button(
            label="📊 Excel",
            data=excel_bytes,
            file_name=f"Orcamento_{prop_formatted}_{cliente.replace(' ', '_') if cliente else 'Cliente'}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

# --- ABA 2: ORÇAMENTOS SALVOS ---
elif menu == "📋 Orçamentos Salvos":
    st.subheader("📋 Orçamentos Salvos no Banco de Dados")
    conn = get_connection()
    c = conn.cursor()
    
    status_list = get_status_list()
    
    c_f1, c_f2 = st.columns([1, 2])
    with c_f1:
        filtro_status = st.selectbox("Filtrar por Status Comercial", ["Todos"] + status_list)
    with c_f2:
        termo_busca = st.text_input("🔍 Pesquisar por Cliente, Nº Proposta ou Consultor", placeholder="Digite o nome, número ou consultor...").strip()

    query = "SELECT id, proposta_num, cliente, consultor, data, total_liquido, status FROM orcamentos WHERE 1=1"
    params = []

    if filtro_status != "Todos":
        query += " AND status = ?"
        params.append(filtro_status)

    if termo_busca:
        query += " AND (cliente LIKE ? OR CAST(proposta_num AS TEXT) LIKE ? OR consultor LIKE ?)"
        busca_param = f"%{termo_busca}%"
        params.extend([busca_param, busca_param, busca_param])

    query += " ORDER BY id DESC"
    
    df_orc = pd.read_sql_query(query, conn, params=params)
    
    if not df_orc.empty:
        for idx, row in df_orc.iterrows():
            with st.container():
                col_a1, col_a2, col_a3, col_a4, col_a5, col_a6, col_a7, col_a8 = st.columns([1.0, 2.0, 1.2, 0.9, 0.9, 0.7, 0.7, 0.7])
                
                p_fmt = f"{int(row['proposta_num']):04d}" if str(row['proposta_num']).isdigit() else str(row['proposta_num'])
                status_atual = row.get('status', status_list[0] if status_list else 'Em Análise')
                
                col_a1.write(f"**#{p_fmt}**")
                col_a2.write(f"**{row['cliente']}**\n\n`{status_atual}`")
                col_a3.write(f"R$ {row['total_liquido']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
                
                if col_a4.button("✏️ Editar", key=f"btn_edit_orc_{row['id']}"):
                    c.execute("SELECT * FROM orcamentos WHERE id = ?", (row['id'],))
                    o = dict(c.fetchone())
                    
                    st.session_state.edit_index = o['id']
                    st.session_state.cli_prop = o['proposta_num']
                    st.session_state.cli_nome = o['cliente'] or ""
                    st.session_state.cli_contato = o['contato'] or ""
                    st.session_state.cli_tipo = o['tipo_contato'] or "Residencial"
                    st.session_state.cli_tel = o['telefone'] or ""
                    st.session_state.cli_email = o['email'] or ""
                    st.session_state.cli_consultor = o['consultor'] or "Sem Consultor"
                    st.session_state.cli_prazo = o['prazo_entrega'] or ""
                    st.session_state.cli_cond = o['condicoes_pagamento'] or ""
                    st.session_state.cli_obs = o['observacoes'] or ""
                    st.session_state.cli_status = o.get('status', status_list[0] if status_list else 'Em Análise')
                    
                    st.session_state.expand_ambientes = True
                    st.session_state.form_version += 1
                    
                    ambs_db = []
                    c.execute("SELECT id, nome_ambiente, especificacoes, total_ambiente FROM ambientes WHERE orcamento_id = ? ORDER BY ordem", (row['id'],))
                    for amb_row in c.fetchall():
                        amb_dict = dict(amb_row)
                        itens_db = []
                        c.execute("SELECT descricao, valor, eh_opcional FROM itens WHERE ambiente_id = ? ORDER BY ordem", (amb_dict['id'],))
                        for item_row in c.fetchall():
                            item_dict = dict(item_row)
                            itens_db.append({
                                'id': str(uuid.uuid4()),
                                'descricao': item_dict['descricao'],
                                'valor': item_dict['valor'],
                                'eh_opcional': bool(item_dict['eh_opcional'])
                            })
                        ambs_db.append({
                            'nome': amb_dict['nome_ambiente'],
                            'especificacoes': amb_dict['especificacoes'] or "",
                            'total_ambiente': amb_dict['total_ambiente'],
                            'itens': itens_db
                        })
                    st.session_state.ambientes = ambs_db
                    st.session_state.change_tab_to = "➕ Novo / Editar Orçamento"
                    st.rerun()

                if col_a5.button("📋 Clonar", key=f"btn_clone_orc_{row['id']}"):
                    c.execute("SELECT * FROM orcamentos WHERE id = ?", (row['id'],))
                    o = dict(c.fetchone())
                    
                    st.session_state.edit_index = None
                    st.session_state.cli_prop = get_proxima_proposta()
                    st.session_state.cli_nome = f"{o['cliente']} (Cópia)"
                    st.session_state.cli_contato = o['contato'] or ""
                    st.session_state.cli_tipo = o['tipo_contato'] or "Residencial"
                    st.session_state.cli_tel = o['telefone'] or ""
                    st.session_state.cli_email = o['email'] or ""
                    st.session_state.cli_consultor = o['consultor'] or "Sem Consultor"
                    st.session_state.cli_prazo = o['prazo_entrega'] or ""
                    st.session_state.cli_cond = o['condicoes_pagamento'] or ""
                    st.session_state.cli_obs = o['observacoes'] or ""
                    st.session_state.cli_status = status_list[0] if status_list else "Em Análise"
                    
                    st.session_state.expand_ambientes = True
                    st.session_state.form_version += 1
                    
                    ambs_db = []
                    c.execute("SELECT id, nome_ambiente, especificacoes, total_ambiente FROM ambientes WHERE orcamento_id = ? ORDER BY ordem", (row['id'],))
                    for amb_row in c.fetchall():
                        amb_dict = dict(amb_row)
                        itens_db = []
                        c.execute("SELECT descricao, valor, eh_opcional FROM itens WHERE ambiente_id = ? ORDER BY ordem", (amb_dict['id'],))
                        for item_row in c.fetchall():
                            item_dict = dict(item_row)
                            itens_db.append({
                                'id': str(uuid.uuid4()),
                                'descricao': item_dict['descricao'],
                                'valor': item_dict['valor'],
                                'eh_opcional': bool(item_dict['eh_opcional'])
                            })
                        ambs_db.append({
                            'nome': amb_dict['nome_ambiente'],
                            'especificacoes': amb_dict['especificacoes'] or "",
                            'total_ambiente': amb_dict['total_ambiente'],
                            'itens': itens_db
                        })
                    st.session_state.ambientes = ambs_db
                    st.session_state.change_tab_to = "➕ Novo / Editar Orçamento"
                    st.rerun()

                c.execute("SELECT * FROM orcamentos WHERE id = ?", (row['id'],))
                o_saved = dict(c.fetchone())
                ambs_saved = []
                c.execute("SELECT id, nome_ambiente, especificacoes, total_ambiente FROM ambientes WHERE orcamento_id = ? ORDER BY ordem", (row['id'],))
                for amb_row in c.fetchall():
                    amb_dict = dict(amb_row)
                    itens_saved = []
                    c.execute("SELECT descricao, valor, eh_opcional FROM itens WHERE ambiente_id = ? ORDER BY ordem", (amb_dict['id'],))
                    for item_row in c.fetchall():
                        item_dict = dict(item_row)
                        itens_saved.append({'descricao': item_dict['descricao'], 'valor': item_dict['valor'], 'eh_opcional': bool(item_dict['eh_opcional'])})
                    ambs_saved.append({'nome': amb_dict['nome_ambiente'], 'especificacoes': amb_dict['especificacoes'] or "", 'total_ambiente': amb_dict['total_ambiente'], 'itens': itens_saved})

                cli_saved_info = {
                    'proposta_num': p_fmt, 'cliente': o_saved['cliente'], 'contato': o_saved['contato'], 'tipo_contato': o_saved['tipo_contato'],
                    'telefone': o_saved['telefone'], 'email': o_saved['email'], 'consultor': o_saved['consultor'], 'data': o_saved['data'],
                    'dias_validade': o_saved['dias_validade'], 'validade': o_saved['validade'], 'prazo_entrega': o_saved['prazo_entrega'],
                    'condicoes_pagamento': o_saved['condicoes_pagamento'], 'observacoes': o_saved['observacoes']
                }

                pdf_saved_bytes = gerar_pdf_orcamento(cli_saved_info, ambs_saved)
                word_saved_bytes = gerar_word_orcamento(cli_saved_info, ambs_saved)
                
                col_a6.download_button(
                    label="📄 PDF",
                    data=pdf_saved_bytes,
                    file_name=f"Orcamento_{p_fmt}_{o_saved['cliente'].replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    key=f"btn_pdf_orc_{row['id']}"
                )

                col_a7.download_button(
                    label="📝 DOCX",
                    data=word_saved_bytes,
                    file_name=f"Orcamento_{p_fmt}_{o_saved['cliente'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"btn_word_orc_{row['id']}"
                )

                if col_a8.button("🗑️ Excluir", key=f"btn_del_orc_{row['id']}"):
                    st.session_state.confirm_del = f"orc_{row['id']}"

                if st.session_state.confirm_del == f"orc_{row['id']}":
                    st.warning(f"⚠️ Confirma a exclusão permanente do Orçamento Proposta #{p_fmt}?")
                    c_del1, c_del2 = st.columns(2)
                    if c_del1.button("✅ Confirmar", key=f"conf_del_orc_{row['id']}"):
                        c.execute("DELETE FROM orcamentos WHERE id = ?", (row['id'],))
                        conn.commit()
                        st.session_state.confirm_del = None
                        st.rerun()
                    if c_del2.button("❌ Cancelar", key=f"canc_del_orc_{row['id']}"):
                        st.session_state.confirm_del = None
                        st.rerun()

                st.markdown("---")
    else:
        st.info("Nenhum orçamento encontrado com os filtros e busca atuais.")

# --- ABA 3: CONFIGURAÇÕES E LOGO ---
elif menu == "⚙️ Configurações":
    st.subheader("⚙️ Configurações da Empresa, Consultores, Status e Segurança")
    conn = get_connection()
    c = conn.cursor()
    
    st.markdown("#### 🏢 Logo e Dados da Fábrica")
    config = get_config()
    
    uploaded_logo = st.file_uploader("Enviar Logo da Empresa (PNG / JPG)", type=["png", "jpg", "jpeg"])
    if uploaded_logo:
        logo_path = os.path.join("logo_empresa.png")
        with open(logo_path, "wb") as f:
            f.write(uploaded_logo.getbuffer())
        c.execute("UPDATE config_empresa SET logo_path = ? WHERE id = 1", (logo_path,))
        conn.commit()
        st.success("Logo atualizada com sucesso!")

    logo_actual = config.get('logo_path', '')
    if logo_actual and os.path.exists(logo_actual):
        st.image(logo_actual, width=200, caption="Logo Atual")

    with st.form("form_config_empresa"):
        nome_empresa = st.text_input("Razão Social", value=config.get('nome_empresa', 'Fábrica de Móveis Laurenti Ltda'))
        cnpj = st.text_input("CNPJ", value=config.get('cnpj', '44.331.015/0001-08'))
        ie = st.text_input("Inscrição Estadual (IE)", value=config.get('ie', '186000158114'))
        endereco = st.text_input("Endereço Completo", value=config.get('endereco', 'Rua Henrique Villa, 59- Jardim Maria Emília. CEP: 15960000, ARIRANHA-SP'))
        telefone = st.text_input("Telefone", value=config.get('telefone', '(17) 3576-1464'))
        email = st.text_input("E-mail", value=config.get('email', 'contato@laurentimoveis.com.br'))
        
        logo_largura = st.slider("Largura da Logo no PDF (cm)", min_value=3.0, max_value=8.0, value=float(config.get('logo_largura', 5.0)), step=0.5)
        
        if st.form_submit_button("💾 Salvar Dados da Empresa"):
            c.execute("""
                UPDATE config_empresa
                SET nome_empresa = ?, cnpj = ?, ie = ?, endereco = ?, telefone = ?, email = ?, logo_largura = ?
                WHERE id = 1
            """, (nome_empresa, cnpj, ie, endereco, telefone, email, logo_largura))
            conn.commit()
            st.success("Dados da empresa salvos!")
            st.rerun()

    st.markdown("---")
    
    st.markdown("#### 🔑 Alterar Senha do Sistema")
    with st.form("form_alterar_senha"):
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            nova_senha = st.text_input("Nova Senha", type="password")
        with col_p2:
            confirma_senha = st.text_input("Confirmar Nova Senha", type="password")
            
        if st.form_submit_button("🔐 Atualizar Senha"):
            if not nova_senha:
                st.error("A senha não pode ser vazia.")
            elif nova_senha != confirma_senha:
                st.error("As senhas digitadas não coincidem.")
            else:
                c.execute("UPDATE config_empresa SET senha_hash = ? WHERE id = 1", (hash_senha(nova_senha),))
                conn.commit()
                st.success("Senha do sistema alterada com sucesso!")

    st.markdown("---")
    col_cfg1, col_cfg2 = st.columns(2)
    
    with col_cfg1:
        st.markdown("#### 🏷️ Gestão de Status Comercial")
        
        with st.form("form_add_status", clear_on_submit=True):
            novo_st = st.text_input("Nome do Novo Status Comercial")
            if st.form_submit_button("➕ Cadastrar Status"):
                if novo_st:
                    try:
                        c.execute("INSERT INTO status_comercial (nome) VALUES (?)", (novo_st.strip(),))
                        conn.commit()
                        st.success(f"Status '{novo_st}' cadastrado!")
                        st.rerun()
                    except sqlite3.IntegrityError:
                        st.error("Este status já está cadastrado.")

        status_atuais = get_status_list()
        for st_item in status_atuais:
            col_s1, col_s2 = st.columns([3, 1])
            col_s1.write(f"🏷️ **{st_item}**")
            
            if col_s2.button("🗑️ Excluir", key=f"del_status_btn_{st_item}"):
                st.session_state.confirm_del = f"status_{st_item}"

            if st.session_state.confirm_del == f"status_{st_item}":
                st.warning(f"⚠️ Confirma excluir o status '{st_item}'?")
                st.info("Orçamentos com este status serão migrados para o status padrão.")
                c_del1, c_del2 = st.columns(2)
                
                if c_del1.button("✅ Confirmar", key=f"conf_del_status_{st_item}"):
                    c.execute("DELETE FROM status_comercial WHERE nome = ?", (st_item,))
                    conn.commit()
                    
                    status_restantes = get_status_list()
                    fallback_status = status_restantes[0] if status_restantes else "Em Análise"
                    if not status_restantes:
                        c.execute("INSERT INTO status_comercial (nome) VALUES (?)", ("Em Análise",))
                        conn.commit()
                        fallback_status = "Em Análise"
                    
                    c.execute("UPDATE orcamentos SET status = ? WHERE status = ?", (fallback_status, st_item))
                    conn.commit()
                    
                    st.session_state.confirm_del = None
                    st.rerun()
                    
                if c_del2.button("❌ Cancelar", key=f"canc_del_status_{st_item}"):
                    st.session_state.confirm_del = None
                    st.rerun()

    with col_cfg2:
        st.markdown("#### 👤 Gestão de Consultores")
        
        with st.form("form_add_consultor", clear_on_submit=True):
            novo_c = st.text_input("Nome do Novo Consultor")
            if st.form_submit_button("➕ Cadastrar Consultor"):
                if novo_c:
                    c.execute("INSERT INTO consultores (nome) VALUES (?)", (novo_c.strip(),))
                    conn.commit()
                    st.rerun()

        df_c = get_consultores()
        if not df_c.empty:
            for _, r in df_c.iterrows():
                col_cons1, col_cons2 = st.columns([3, 1])
                col_cons1.write(f"👤 **{r['nome']}**")
                
                if col_cons2.button("🗑️ Excluir", key=f"del_cons_btn_{r['id']}"):
                    st.session_state.confirm_del = f"cons_{r['id']}"

                if st.session_state.confirm_del == f"cons_{r['id']}":
                    st.warning(f"⚠️ Remover o consultor '{r['nome']}'?")
                    c_del1, c_del2 = st.columns(2)
                    if c_del1.button("✅ Confirmar", key=f"conf_del_cons_{r['id']}"):
                        c.execute("DELETE FROM consultores WHERE id = ?", (r['id'],))
                        conn.commit()
                        st.session_state.confirm_del = None
                        st.rerun()
                    if c_del2.button("❌ Cancelar", key=f"canc_del_cons_{r['id']}"):
                        st.session_state.confirm_del = None
                        st.rerun()